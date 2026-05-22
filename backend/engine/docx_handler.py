"""
DOCX file handler — extract text with run-level metadata and apply replacements.

Handles: body paragraphs, tables, headers, footers.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from lxml import etree
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .cross_run import replace_entities_in_paragraph_docx, replace_entities_in_text


# TODO: Images containing text (would need OCR)
# TODO: Embedded Excel objects
# TODO: Charts and SmartArt


@dataclass
class TextBlock:
    """A block of text extracted from a DOCX file with its location."""
    section: str  # e.g. "Paragraphe 3", "Tableau 1 - Cellule (2,3)", "En-tête"
    text: str
    paragraph_ref: object = field(default=None, repr=False)


def extract_text_blocks(file_bytes: bytes) -> tuple[Document, list[TextBlock]]:
    """Extract all text blocks from a DOCX file.

    Returns:
        doc: the Document object (for later modification)
        blocks: list of TextBlock with text and location metadata
    """
    doc = Document(io.BytesIO(file_bytes))
    blocks = []

    # Body paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            blocks.append(TextBlock(
                section=f"Paragraphe {i + 1}",
                text=text,
                paragraph_ref=para,
            ))

    # Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if text:
                        blocks.append(TextBlock(
                            section=f"Tableau {t_idx + 1} — Cellule ({r_idx + 1},{c_idx + 1})",
                            text=text,
                            paragraph_ref=para,
                        ))

    # Headers and footers
    for section in doc.sections:
        for header_type, label in [
            (section.header, "En-tête"),
            (section.footer, "Pied de page"),
        ]:
            if header_type and header_type.is_linked_to_previous is False or header_type.paragraphs:
                for para in header_type.paragraphs:
                    text = para.text.strip()
                    if text:
                        blocks.append(TextBlock(
                            section=label,
                            text=text,
                            paragraph_ref=para,
                        ))

    return doc, blocks


def apply_replacements(
    doc: Document,
    entities_sorted: list[tuple[str, str]],
    alt_mapping: dict[str, str] | None = None,
):
    """Apply entity replacements across the entire DOCX document.

    Args:
        doc: python-docx Document object
        entities_sorted: list of (original, placeholder) sorted by len DESC
        alt_mapping: dict of original_alt_text -> placeholder for alt text replacement
    """
    # Body paragraphs
    for para in doc.paragraphs:
        replace_entities_in_paragraph_docx(para, entities_sorted)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_entities_in_paragraph_docx(para, entities_sorted)

    # Headers and footers
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer:
                for para in header_footer.paragraphs:
                    replace_entities_in_paragraph_docx(para, entities_sorted)

    # Replace alt texts on images
    if alt_mapping:
        _replace_alt_texts_docx(doc, alt_mapping)

    # Anonymize metadata
    _clean_metadata(doc, entities_sorted)


def extract_alt_texts(doc: Document) -> list[tuple[str, str]]:
    """Extract all non-empty alt texts from images in a DOCX.

    Looks for 'descr' attributes on docPr elements inside drawings.

    Returns:
        list of (alt_text, location_label)
    """
    alt_texts: list[tuple[str, str]] = []
    WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

    def _scan_paragraphs(paragraphs, label: str):
        for para in paragraphs:
            for drawing in para._element.iter():
                if etree.QName(drawing).localname == "docPr":
                    descr = (drawing.get("descr") or "").strip()
                    if descr:
                        alt_texts.append((descr, label))

    _scan_paragraphs(doc.paragraphs, "Corps du document")

    for t_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                _scan_paragraphs(cell.paragraphs, f"Tableau {t_idx + 1}")

    for section in doc.sections:
        if section.header:
            _scan_paragraphs(section.header.paragraphs, "En-tête")
        if section.footer:
            _scan_paragraphs(section.footer.paragraphs, "Pied de page")

    return alt_texts


def _replace_alt_texts_docx(doc: Document, alt_mapping: dict[str, str]):
    """Replace alt texts (descr attributes on docPr) across the DOCX."""
    def _replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            for elem in para._element.iter():
                if etree.QName(elem).localname == "docPr":
                    descr = (elem.get("descr") or "").strip()
                    if descr and descr in alt_mapping:
                        elem.set("descr", alt_mapping[descr])

    _replace_in_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs)

    for section in doc.sections:
        if section.header:
            _replace_in_paragraphs(section.header.paragraphs)
        if section.footer:
            _replace_in_paragraphs(section.footer.paragraphs)


def _clean_metadata(doc: Document, entities_sorted: list[tuple[str, str]] | None = None):
    """Anonymize/clear metadata from document properties.

    The title is anonymized via entity replacement; other fields are cleared.
    """
    props = doc.core_properties

    # Anonymize title with entity replacement instead of just clearing
    if entities_sorted:
        try:
            title = props.title or ""
            if title:
                props.title = replace_entities_in_text(title, entities_sorted)
        except Exception:
            try:
                props.title = ""
            except Exception:
                pass
    else:
        try:
            props.title = ""
        except Exception:
            pass

    try:
        props.author = ""
    except Exception:
        pass
    try:
        props.last_modified_by = ""
    except Exception:
        pass
    try:
        props.comments = ""
    except Exception:
        pass
    try:
        props.keywords = ""
    except Exception:
        pass
    try:
        props.subject = ""
    except Exception:
        pass
    # company is stored in extended properties, not always accessible via python-docx
    # but we try
    try:
        props.company = ""
    except (AttributeError, Exception):
        pass


def save_to_bytes(doc: Document) -> bytes:
    """Save a Document to bytes."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
